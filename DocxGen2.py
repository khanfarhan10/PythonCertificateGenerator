"""
import re
from docx import Document


def docx_replace_regex(doc_obj, regex, replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


regex1 = re.compile(r"YOUR_AWESOME_NAME")
replace1 = r"Farhan Hai Khan"
filename = "templates/1.docx"
doc = Document(filename)
#docx_replace_regex(doc, regex1, replace1)
doc.save('outputs/generated_doc.docx')
"""

from docx import Document

document = Document('templates/1.docx')

dic = {'YOUR_AWESOME_NAME': 'Farhan Hai Khan'}
for p in document.paragraphs:
    inline = p.runs
    for i in range(len(inline)):
        text = inline[i].text
        if text in dic.keys():
            text = text.replace(text, dic[text])
            inline[i].text = text

document.save('outputs/generated_doc.docx')
